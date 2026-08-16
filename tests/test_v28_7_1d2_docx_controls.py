from io import BytesIO
import zipfile

from docx import Document
from lxml import etree

import project_domain_normalization as pdn
from docx_control_text import extract_docx_paragraphs_preserving_controls
from file_analyst import extract_evidence_units
from project_batch_ingestion import _docx_text

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W14 = "http://schemas.microsoft.com/office/word/2010/wordml"


def _docx_with_nested_checkbox_controls(*, checked_yes: bool, checked_no: bool) -> bytes:
    doc = Document()
    paragraph = doc.add_paragraph()
    paragraph.add_run("CONCORRENCIA: SIM, Quais agências: NÃO")
    buf = BytesIO()
    doc.save(buf)

    source = zipfile.ZipFile(BytesIO(buf.getvalue()))
    xml = source.read("word/document.xml")
    root = etree.fromstring(xml)
    ns = {"w": W}
    p = root.xpath("//w:body/w:p", namespaces=ns)[0]
    for child in list(p):
        p.remove(child)

    def add_run(text: str) -> None:
        r = etree.SubElement(p, f"{{{W}}}r")
        t = etree.SubElement(r, f"{{{W}}}t")
        t.text = text

    def add_checkbox(checked: bool) -> None:
        sdt = etree.SubElement(p, f"{{{W}}}sdt")
        pr = etree.SubElement(sdt, f"{{{W}}}sdtPr")
        checkbox = etree.SubElement(pr, f"{{{W14}}}checkbox")
        state = etree.SubElement(checkbox, f"{{{W14}}}checked")
        state.set(f"{{{W14}}}val", "1" if checked else "0")
        content = etree.SubElement(sdt, f"{{{W}}}sdtContent")
        r = etree.SubElement(content, f"{{{W}}}r")
        t = etree.SubElement(r, f"{{{W}}}t")
        t.text = "☒" if checked else "☐"

    add_run("CONCORRENCIA: ")
    add_checkbox(checked_yes)
    add_run("SIM, Quais agências: ")
    add_checkbox(checked_no)
    add_run("NÃO")

    new_xml = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")
    out = BytesIO()
    with zipfile.ZipFile(out, "w") as target:
        for item in source.infolist():
            target.writestr(item, new_xml if item.filename == "word/document.xml" else source.read(item.filename))
    source.close()
    return out.getvalue()


def test_python_docx_control_extractor_preserves_visible_checkbox_state():
    payload = _docx_with_nested_checkbox_controls(checked_yes=False, checked_no=True)
    # This is the exact historical failure mode: Paragraph.text drops nested sdt text.
    plain = Document(BytesIO(payload)).paragraphs[0].text
    assert plain == "CONCORRENCIA: SIM, Quais agências: NÃO"

    paragraphs = extract_docx_paragraphs_preserving_controls(payload)
    assert paragraphs[0]["text"] == "CONCORRENCIA: ☐SIM, Quais agências: ☒NÃO"


def test_file_analyst_and_batch_text_preserve_controls_for_future_ingests():
    payload = _docx_with_nested_checkbox_controls(checked_yes=False, checked_no=True)
    units = extract_evidence_units("briefing.docx", payload)
    assert units[0].content_text == "CONCORRENCIA: ☐SIM, Quais agências: ☒NÃO"
    assert "CONCORRENCIA: ☐SIM, Quais agências: ☒NÃO" in _docx_text(payload)


def test_domain_normalization_recovers_old_lossy_evidence_from_stored_docx(monkeypatch):
    payload = _docx_with_nested_checkbox_controls(checked_yes=False, checked_no=True)
    monkeypatch.setattr(pdn, "get_bytes", lambda client, bucket_name, path: payload)

    match = pdn._direct_commercial_process_evidence(
        briefing_documents=[{"content_sha256": "sha", "storage_bucket": "r2:test", "storage_path": "brief.docx"}],
        asset_by_sha={"sha": {"id": "asset", "storage_bucket": "r2:test", "storage_path": "brief.docx"}},
        evidence_by_asset={"asset": [{
            "id": "old-eu",
            "unit_type": "paragraph",
            "ordinal": 7,
            "locator": {"paragraph_index": 1},
            "content_text": "CONCORRENCIA: \tSIM, Quais agências: \tNÃO",
        }]},
        client=object(),
    )
    assert match and match["id"] == "old-eu"
    assert match["_source_semantic_text"] == "CONCORRENCIA: ☐SIM, Quais agências: ☒NÃO"
    assert match["_semantic_recovery_method"] == "stored_docx_ooxml_controls"


def test_domain_normalization_does_not_infer_direct_when_yes_is_checked(monkeypatch):
    payload = _docx_with_nested_checkbox_controls(checked_yes=True, checked_no=False)
    monkeypatch.setattr(pdn, "get_bytes", lambda client, bucket_name, path: payload)

    match = pdn._direct_commercial_process_evidence(
        briefing_documents=[{"content_sha256": "sha", "storage_bucket": "r2:test", "storage_path": "brief.docx"}],
        asset_by_sha={"sha": {"id": "asset", "storage_bucket": "r2:test", "storage_path": "brief.docx"}},
        evidence_by_asset={"asset": [{
            "id": "old-eu",
            "unit_type": "paragraph",
            "ordinal": 7,
            "locator": {"paragraph_index": 1},
            "content_text": "CONCORRENCIA: \tSIM, Quais agências: \tNÃO",
        }]},
        client=object(),
    )
    assert match is None
